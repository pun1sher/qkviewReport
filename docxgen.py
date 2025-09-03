from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

class docxgen:
    
    """Generates Word documents from parsed data"""
    
    def __init__(self, output_path: str):
        self.output_path = Path(output_path)
        self.output_path.mkdir(parents=True, exist_ok=True)
    
    def create_device_report(self, device_info, license_info, cfg_interfaces, object_counts, graph_objects, diags) -> str:
        """Create comprehensive device report"""
        hostName = device_info[0][1]
        qkviewDate = device_info[7][1]
        qkvDate = datetime.strftime(datetime.strptime(qkviewDate, "%Y-%m-%d %H:%M:%S"), "%Y-%m-%d")

        document = Document()
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = f'Qkview Report\t\t{qkvDate}'
        document.add_heading(hostName, 0)
        
        # Add device information table
        self._add_device_info_table(document, device_info)

        # Add license information table
        self._add_license_info_table(document,license_info)
        
        # Add configured Interface table
        self._add_config_interface_table(document,cfg_interfaces)
        # Add object counts table
        self._add_object_counts_table(document, object_counts)
        
        # Add Performance Graphs
        self._add_performance_graphs(document, graph_objects)

        # Add Diagnostic report (Critical, High and Medium)
        self._add_diagnostic_table(document, diags)

        # Add VS statistics table
       # self._add_vs_stats_table(document, vs_stats)
        
        # Save document
        
        filename = f"{hostName}_{qkvDate}.docx".lower()
        filepath = self.output_path / filename
        document.save(str(filepath))
        
        return str(filepath)
    
    def _add_device_info_table(self, document, device_info):
        """Add device information table to document"""
        document.add_heading('Device Information', level=2)
        table = document.add_table(rows=9, cols=2, style='TableGrid')
        
        for i, (key, value) in enumerate(device_info):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value

    def _add_license_info_table(self, document, license_info):
        document.add_heading('Licensing', level=2)
        numRows = len(license_info)
        table = document.add_table(rows=numRows, cols=2, style='TableGrid')
        for i,(key,value) in enumerate(license_info):
            table.cell(i,0).text = key
            table.cell(i,1).text = value

    def _add_config_interface_table(self, document, cfg_interfaces):
        headerValues = ['Inf Name', 'Status', 'Media', 'Trunk']
        if len(cfg_interfaces) > 0:
            document.add_heading('Configured Interfaces', level=2)
            numRows = len(cfg_interfaces ) + 1
            numCols = 4
            table = document.add_table(rows=numRows, cols=numCols, style = 'TableGrid')
            for i, headerName in enumerate(headerValues):
                table.cell(0,i).text = headerName
            row = 1
            for key in cfg_interfaces.keys():
                table.cell(row,0).text = key
                table.cell(row,1).text = cfg_interfaces[key]['status']
                table.cell(row,2).text = cfg_interfaces[key]['media']
                table.cell(row,3).text = cfg_interfaces[key]['trunk']
                row += 1


    def _add_diagnostic_table(self,document,diags):
        headerValues = ['Name', 'CVE Number', 'KB Article', 'KB Link', 'Fixed In Version', 'Severity']
        section = document.add_section()
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        document.add_heading('Diagnostics', level=2)
        numRows = len(diags) + 1
        table = document.add_table(rows=numRows, cols=6, style='TableGrid')
        for i,(headerName) in enumerate(headerValues):
            table.cell(0,i).text = headerName   
        row = 1
        for (name, kbArticle, kbLink, fixedVersion, importance, cveNum) in diags:
            table.cell(row,0).text = name
            table.cell(row,1).text = cveNum
            table.cell(row,2).text = kbArticle
            table.cell(row,3).text = kbLink
            table.cell(row,4).text = fixedVersion
            table.cell(row,5).text = importance
            row += 1

    def _add_object_counts_table(self,document, object_counts):

        document.add_heading('Configured Object Counts', level=2)

        numRows = len(object_counts)
        table = document.add_table(rows=numRows,cols=2, style='TableGrid')
        for i, (key,value) in enumerate(object_counts):
            table.cell(i,0).text = key
            table.cell(i,1).text = str(value)

    def _add_performance_graphs(self, document, graph_objects):
        
        document.add_page_break()
        document.add_heading('Performance Graphs', level=2)
        for img in graph_objects:
            if 'active_conn' in img:
                document.add_heading('Active Connections', level=3)
            elif 'by_core' in img:
                document.add_heading('CPU Usage By Core', level=3)
            elif 'system' in img:
                document.add_heading('System CPU Usage', level=3)
            elif 'memory' in img:
                document.add_heading('Memory Breakdown', level=3)
            elif 'new_connections' in img:
                document.add_heading('New Connections', level=3)
            elif 'throughput' in img:
                document.add_heading('Throughput (bits)', level=3)
             
            document.add_picture( img, width=Inches(6.2))

    